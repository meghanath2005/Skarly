"""Build the illustrated Skarly research-internship report as a polished DOCX."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "docs" / "research-report"
OUTPUT = REPORT_DIR / "Skarly_Research_Internship_Report.docx"
ARTIFACTS = ROOT / "research" / "artifacts"
SCREENSHOTS = ROOT / "docs" / "ui-screenshots"
WORDMARK = ROOT / "lyricmorph-mobile" / "assets" / "skarly-wordmark.png"

GOLD = "C6AA6A"
GOLD_LIGHT = "F4E9CD"
NAVY = "0B1B2B"
INK = "17202A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6570"
PALE = "F4F6F9"
GREEN = "2E6F58"
RED = "9B1C1C"


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def set_paragraph_shading(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=9, color=MUTED, italic=True)
    return p


def add_figure(doc, path, caption, width=6.15):
    path = Path(path)
    if not path.is_file():
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)
    return True


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True, color=INK)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, color=INK)
    else:
        run = p.add_run(text)
        set_font(run, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text), color=INK)
    return p


def add_numbered(doc, title, body):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(title + ". "), bold=True, color=INK)
    set_font(p.add_run(body), color=INK)
    return p


def add_callout(doc, label, text, fill=PALE, accent=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.1
    set_paragraph_shading(p, fill, accent)
    set_font(p.add_run(label.upper() + "  "), size=9, bold=True, color=accent)
    set_font(p.add_run(text), size=10.5, color=INK)
    return p


def add_source_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), size=8.5, color=MUTED, italic=True)
    return p


def add_data_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    mark_repeat_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=9.5, color="FFFFFF", bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for col_index, (cell, value) in enumerate(zip(cells, values)):
            if row_index % 2:
                set_cell_shading(cell, "F8F9FB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(str(value)), size=9.2, color=INK, bold=(col_index == 0))
        set_table_geometry(table, widths_dxa)
    return table


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    set_paragraph_shading(p, "111318")
    for index, line in enumerate(lines):
        run = p.add_run(line + ("\n" if index < len(lines) - 1 else ""))
        set_font(run, name="Consolas", size=8.5, color="E9EDF2")
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)


def configure_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run("SKARLY  /  RESEARCH INTERNSHIP REPORT"), size=8, color=MUTED, bold=True)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_after = Pt(0)
        set_font(fp.add_run("Skarly research artifact  |  "), size=8, color=MUTED)
        add_page_field(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(16)
    set_font(p.add_run("RESEARCH INTERNSHIP PROJECT"), size=10, color=GOLD, bold=True)
    if WORDMARK.is_file():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_after = Pt(12)
        pic.add_run().add_picture(str(WORDMARK), width=Inches(5.6))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("Audio Intelligence and Five-Version Music Generation"), size=25, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    set_font(subtitle.add_run("System design, neural models, signal analysis, UI evidence, and reproducible procedure"), size=13, color=MUTED, italic=True)
    add_callout(doc, "Project boundary", "The working vocal-to-music and music-to-music paths were tested as a guest and left unchanged. New work is limited to a result-screen player, research notebook, screenshots, documentation, and packaging.", fill=GOLD_LIGHT, accent=GOLD)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(26)
    set_font(meta.add_run("Prepared 18 July 2026\nSkarly local research build\nCase-study audio: 225.048-second vocal recording"), size=10.5, color=MUTED)
    doc.add_page_break()


def build_report():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    summary = json.loads((ARTIFACTS / "research_summary.json").read_text(encoding="utf-8"))
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    add_cover(doc)

    doc.add_heading("Executive abstract", level=1)
    add_body(doc, "Skarly is a local full-stack creator studio that accepts a vocal, full song, or music reference; analyses the complete decoded recording; creates five differentiated producer blueprints; generates new backing arrangements with ACE-Step; preserves or separates the singer according to the selected mode; and exports playable mixes, stems, maps, and disclosure metadata.")
    add_callout(doc, "Outcome", "The live guest test completed successfully with the supplied MP3: private upload, full-song analysis, five producer plans, CUDA generation, five playable versions, and export preparation all rendered in the browser. The result UI now includes a compact message-style player beneath the five final versions.", fill="EAF4EF", accent=GREEN)
    add_data_table(doc, ["Evidence", "Observed result"], [
        ["Input duration", "225.048 seconds, complete-song scope"],
        ["Signal map", "23 vocal phrases across 6 detected sections"],
        ["Production estimate", "Approximately 170 BPM; C# minor; up-tempo pop start"],
        ["Runtime", "NVIDIA GeForce RTX 5070 Laptop GPU; ACE-Step 1.5 turbo"],
        ["Generation telemetry", "7,718 MB peak VRAM; 113-second render"],
        ["Diversity gate", "10/10 producer pairs passed; thresholds remain prototype"],
    ], [2200, 7160])
    add_source_line(doc, "Source: live local guest run and executed notebook, 18 July 2026.")

    doc.add_heading("1. Scope, originality, and safety boundary", level=1)
    add_body(doc, "This research pass treats the existing application as the system under study. The implementation was mapped from the repository's own modules and run-time evidence. The notebook's signal-processing code was written specifically for this project using transparent NumPy/SciPy operations. No external template code was inserted into the model, generator, source separation, mixing, or API paths.")
    add_bullet(doc, "Protected working scope: vocal-to-music and music-to-music generation, source preparation, ACE-Step calls, mixing, quality validation, and storage contracts.")
    add_bullet(doc, "Allowed presentation scope: one selected-version player, screenshots, research charts, report assets, and plug-and-play setup documentation.")
    add_bullet(doc, "Data boundary: the attached audio is used for the live test and notebook execution, but it is excluded from Git and is not treated as training data.")
    add_bullet(doc, "Research boundary: checkpoint metrics are reported as observed; limitations and prototype thresholds remain explicit.")

    doc.add_heading("2. Working system architecture", level=1)
    add_figure(doc, ARTIFACTS / "model_architecture.png", "Figure 1. Working architecture and the boundary between analysis, generation, mixing, and export.")
    add_body(doc, "The architecture separates recognition from synthesis. The classifier heads guide routing and producer-plan construction; they do not generate audio and do not clone a singer. ACE-Step generates instrumental backing directions. The adaptive mixer then reuses the preserved vocal and makes space with phrase-aware multiband ducking.")

    doc.add_heading("3. End-to-end procedure", level=1)
    for title, body in [
        ("Local upload", "The guest client transfers the selected audio to local storage and verifies that it exists before analysis."),
        ("Source preparation", "Vocal-only input is preserved. Full-song or music input is routed through Demucs when the selected mode requires separation."),
        ("Complete-song analysis", "The backend decodes the full duration, maps tempo, key, phrases, energy and structure, and aggregates learned predictions over contiguous windows."),
        ("Creator confirmation", "Language, style, BPM, key, mood, and mix focus remain visible. Weak genre evidence does not silently override the creator."),
        ("Five producer plans", "Five distinct palettes specify instrumentation, groove, bass movement, energy arc, stereo treatment, and a generation prompt."),
        ("CUDA generation", "ACE-Step 1.5 renders a new backing for each producer plan and records device, model, VRAM, duration, seed, and fallback state."),
        ("Protected vocal mix", "The original vocal is reused; timing and decoded duration are treated as the source of truth; ducking avoids masking the singer."),
        ("Quality gates", "Decode, duration, loudness, silence, vocal leakage, musical compatibility, and pairwise diversity are checked before presentation."),
        ("Selection and revision", "All five versions stay independently playable. Stem remixing avoids regeneration; single-producer revision preserves the other four."),
        ("Export", "The project prepares WAV, MP3, instrumental, vocal, song map, analysis, seeds, telemetry, disclosure, and a studio bundle."),
    ]:
        add_numbered(doc, title, body)

    doc.add_page_break()
    doc.add_heading("4. Models and algorithms used", level=1)
    add_data_table(doc, ["Component", "Role", "Implementation in this project"], [
        ["ACE-Step 1.5", "Music generation", "Creates five new instrumental backings from producer blueprints on CUDA."],
        ["AutoencoderOobleck VAE", "Shared audio encoder", "Frozen ACE-Step encoder; window latent mean and standard deviation form a 128-D embedding."],
        ["Multi-head NN", "Audio intelligence", "LayerNorm → Linear 256 → GELU → Dropout → Linear 256 → calibrated independent heads."],
        ["Legacy mel-CNN", "Research/backwards compatibility", "64-band log-mel image; Conv2D 24/48/96; global pooling; 96-D embedding; language/genre heads."],
        ["Demucs", "Source separation", "Produces validated vocal or instrumental stems for full-song/music routing."],
        ["Basic Pitch", "Melody/MIDI", "Preferred MIDI extraction; skipped for long inputs when configured thresholds require."],
        ["Whisper", "Lyrics/language evidence", "Optional transcription and language evidence used by vocal analysis."],
        ["Librosa/SciPy features", "Signal evidence", "Tempo, chroma, pitch contour, energy, onset and full-song structure support."],
        ["Adaptive mixer", "Vocal-preserving mix", "Phrase-aware ducking, balance control, exact-duration handling and export validation."],
    ], [1900, 2200, 5260])
    add_source_line(doc, "Source: lyricmorph-backend/training/audio_intelligence.py, training/train_audio_classifier.py, app/worker.py, and app/services/.")

    doc.add_heading("5. Neural network and CNN research", level=1)
    add_figure(doc, ARTIFACTS / "nn_cnn_comparison.png", "Figure 2. Legacy convolutional network compared with the current shared-encoder neural architecture.")
    add_body(doc, "The legacy network is useful pedagogically because its convolution kernels operate directly on local time-frequency shapes. The current direction is better aligned with the application: a pretrained music-audio encoder provides a richer representation, while only small task heads are trained. Frozen upstream weights also reduce the amount of project-specific data needed to learn useful routing decisions.")
    add_figure(doc, ARTIFACTS / "cnn_training_curves.png", "Figure 3. Real training and grouped-validation curves from the retained training-history export.")
    add_callout(doc, "Interpretation", "The selected legacy checkpoint reached 96.77% grouped-validation language accuracy and 54.35% broad-genre accuracy at epoch 15. The language result supports guarded routing; the genre result does not justify silent fine-grained style selection, so creator confirmation remains part of the product.", fill="FFF7E6", accent=GOLD)

    doc.add_heading("6. Case-study audio analysis", level=1)
    add_body(doc, f"The notebook analysed the supplied {summary['duration_seconds']:.3f}-second recording without adding it to the repository. FFmpeg decoded a temporary mono analysis stream; SciPy/NumPy code then computed waveform, RMS, zero-crossing activity, STFT, mel energy, MFCCs, spectral flux, chroma, and a transparent key hypothesis.")
    add_figure(doc, ARTIFACTS / "audio_overview.png", "Figure 4. Waveform, short-time RMS, and zero-crossing activity across the complete recording.")
    add_figure(doc, ARTIFACTS / "spectrogram.png", "Figure 5. Log-frequency STFT spectrogram showing time-varying harmonic and transient content.")
    add_figure(doc, ARTIFACTS / "mel_mfcc.png", "Figure 6. Mel spectrogram and the first 20 MFCCs used to explain CNN-style features.")
    add_figure(doc, ARTIFACTS / "chroma_key.png", "Figure 7. Spectral-flux rhythm evidence and pitch-class energy for an interpretable notebook hypothesis.")
    add_callout(doc, "Why estimates differ", f"The transparent notebook heuristic proposed {summary['notebook_tempo_bpm']:.1f} BPM and {summary['notebook_key_hypothesis']}; the production pipeline reported approximately 170 BPM and C# minor. Different windowing, half/double-time choices, pitch weighting, and confidence logic can produce different hypotheses. The application therefore exposes correction and confirmation rather than presenting one estimate as ground truth.", fill="FDEEEE", accent=RED)

    doc.add_heading("7. Live guest UI validation", level=1)
    add_body(doc, "The browser test used guest mode and the supplied audio through the normal UI. No API was bypassed for the user journey. Upload, analysis, producer selection, processing, version playback, library, profile, and export states were captured as PNG evidence in docs/ui-screenshots/.")
    for filename, caption in [
        ("02-creator-setup.png", "Figure 8. Local creator setup."),
        ("04-upload-empty.png", "Figure 9. Local audio input and mode selection."),
        ("06-audio-detected.png", "Figure 10. Complete-song detection evidence, including language, mood, tempo, timing, key, and readiness."),
        ("08-processing.png", "Figure 11. CUDA generation progress with five-stage status and device telemetry."),
    ]:
        add_figure(doc, SCREENSHOTS / filename, caption, width=6.0)

    doc.add_heading("8. Five final versions and UI enhancement", level=1)
    add_data_table(doc, ["Version", "Producer direction", "Key differentiator"], [
        ["1", "Bollywood Acoustic", "Acoustic guitar, piano, light tabla, warm bass, hook strings"],
        ["2", "Modern Bollywood Pop", "Electronic drums, synth bass, wide pads, plucked hook"],
        ["3", "Sufi Live", "Harmonium, tabla/dholak, claps, live bass, sarangi"],
        ["4", "Punjabi Rhythm", "Dhol, tumbi, punch bass, hand percussion, bright synth accents"],
        ["5", "Cinematic Urban", "Felt piano, atmosphere, deep percussion, strings, sub bass"],
    ], [1000, 2500, 5860])
    add_body(doc, "A new final-version preview sits beneath the five version cards. Its layout follows the supplied voice-message reference: circular headphone avatar, play/pause control, blue progress line and knob, elapsed/total time, rounded pale-green surface, and completion ticks. The component reads the existing selected URL, playback position, duration, and play callback; it introduces no generation or remix behavior.")
    if (SCREENSHOTS / "19-final-preview-added.png").is_file():
        add_figure(doc, SCREENSHOTS / "19-final-preview-added.png", "Figure 12. Reference-inspired final-version player added below the five completed versions.", width=6.0)
    elif (SCREENSHOTS / "10-final-version-playing.png").is_file():
        add_figure(doc, SCREENSHOTS / "10-final-version-playing.png", "Figure 12. Final-version playback confirmed in the live UI.", width=6.0)

    doc.add_heading("9. Quality assurance, ethics, and limitations", level=1)
    add_bullet(doc, "The existing TypeScript application passes `tsc --noEmit` after the UI addition.")
    add_bullet(doc, "The notebook contains 9 executed code cells and zero error outputs; its plots and summary JSON were regenerated from the attached input.")
    add_bullet(doc, "The source audio is excluded from Git and not used for training; creator training contribution remains opt-in and off by default.")
    add_bullet(doc, "No voice cloning path is present in this research work. The singer is preserved as an input stem, not re-synthesized as an identity model.")
    add_bullet(doc, "The diversity gate passed all ten pairs in the case study, but the threshold is explicitly labelled prototype until human calibration requirements are met.")
    add_bullet(doc, "One song is not enough for a production evaluation. Claims require singer-disjoint, rights-cleared datasets and independent human listening panels.")

    doc.add_heading("10. Plug-and-play reproduction", level=1)
    add_body(doc, "The repository is organized so another developer can clone it, create environment files from the checked-in offline examples, start the services in order, open the web UI, and optionally rerun the research notebook with their own permitted audio.")
    add_code_block(doc, [
        "# 1) Start ACE-Step (GPU model service)",
        "powershell -ExecutionPolicy Bypass -File .\\tools\\start-ace-step-api.ps1",
        "",
        "# 2) Start FastAPI backend",
        "powershell -ExecutionPolicy Bypass -File .\\tools\\start-local-studio.ps1",
        "",
        "# 3) Start Expo web UI",
        "cd .\\lyricmorph-mobile",
        "npm install",
        "npm run web",
        "",
        "# 4) Optional research notebook",
        "$env:SKARLY_AUDIO_PATH='C:\\path\\to\\permitted-audio.mp3'",
        "jupyter notebook .\\research\\Skarly_Audio_Intelligence_Research.ipynb",
    ])
    add_callout(doc, "First-run expectation", "ACE-Step weights, GPU drivers, FFmpeg, and local tool paths are environment-specific. The offline example files document the minimum variables without committing secrets.", fill=GOLD_LIGHT, accent=GOLD)

    doc.add_heading("11. Conclusion", level=1)
    add_body(doc, "The project now reads as a research internship build rather than only a product demo: the complete workflow is evidenced, the neural and convolutional paths are explained, real checkpoint curves are plotted, a permitted audio case study is reproducible, the user interface is catalogued, and the working generator boundaries remain intact. The strongest technical design choice is separation of concerns: analysis informs routing, ACE-Step generates new backings, the original singer is protected in mixing, and quality gates stay explicit about what is and is not production-calibrated.")

    doc.add_page_break()
    doc.add_heading("Appendix A. UI screenshot catalogue", level=1)
    add_body(doc, "The following viewport captures document the guest journey and major product states. Full-resolution PNG files are delivered separately in docs/ui-screenshots/.")
    screenshot_files = sorted(SCREENSHOTS.glob("*.png"))
    for index, path in enumerate(screenshot_files, start=1):
        add_figure(doc, path, f"UI-{index:02d}. {path.stem.replace('-', ' ').title()}.", width=5.8)
        if index % 2 == 0 and index < len(screenshot_files):
            doc.add_page_break()

    doc.add_page_break()
    doc.add_heading("Appendix B. Code provenance map", level=1)
    add_data_table(doc, ["Area", "Repository source"], [
        ["Expo UI and result player", "lyricmorph-mobile/App.tsx"],
        ["API contracts and V2 orchestration", "lyricmorph-backend/app/main.py"],
        ["Audio worker and generation routing", "lyricmorph-backend/app/worker.py"],
        ["Frozen encoder and multi-head NN", "lyricmorph-backend/training/audio_intelligence.py"],
        ["Legacy mel-CNN", "lyricmorph-backend/training/train_audio_classifier.py"],
        ["Vocal analysis", "lyricmorph-backend/app/services/vocal_analysis.py"],
        ["ACE-Step wrapper", "lyricmorph-backend/app/services/ace_step_wrapper.py"],
        ["Stem separation", "lyricmorph-backend/app/services/stems_service.py"],
        ["Adaptive mixing", "lyricmorph-backend/app/mixer.py"],
        ["Executable research", "research/Skarly_Audio_Intelligence_Research.ipynb"],
    ], [2600, 6760])
    add_source_line(doc, "All paths are relative to the repository root. Report generated from the local working tree.")

    configure_header_footer(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_report()
